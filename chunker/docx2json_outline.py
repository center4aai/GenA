from __future__ import annotations
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple, Callable
from pathlib import Path
import json, re, statistics

from markitdown import MarkItDown          # DOCX → Markdown  (keeps headings/lists etc.)
# Docs: https://github.com/microsoft/markitdown
from docx import Document                   # direct DOCX parsing
# Docs: https://python-docx.readthedocs.io/
from lxml import etree                      # to read w:outlineLvl, w:numPr/ilvl etc.

WNS = {"w":"http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# ---------- 1) Markdown → дерево (универсально для любого входа) ----------

_heading_atx = re.compile(r'^(?P<h>#{1,6})\s+(?P<title>.+?)(\s+#+\s*)?$', re.UNICODE)
_setext = re.compile(r'^(?P<u>=+|-+)\s*$')
# Регулярное выражение для ссылок markdown: [текст](url)
_link_pattern = re.compile(r'\[([^\]]+)\]\(([^\)]+)\)')

def _iter_markdown_blocks(md: str):
    """Итератор по блокам markdown, сохраняющий ссылки в формате [текст](url)"""
    lines = md.splitlines(); i=0; n=len(lines)
    while i<n:
        line = lines[i]
        m=_heading_atx.match(line)
        if m:
            # Сохраняем ссылки в заголовках
            title = m.group("title").strip()
            yield ("heading", len(m.group("h")), title); i+=1; continue
        if line.strip() and i+1<n:
            m2=_setext.match(lines[i+1])
            if m2:
                # Сохраняем ссылки в setext заголовках
                title = line.strip()
                yield ("heading", 1 if m2.group("u").startswith("=") else 2, title); i+=2; continue
        buf=[line]; i+=1
        while i<n and lines[i].strip()!="":
            if _heading_atx.match(lines[i]): break
            if i+1<n and lines[i].strip() and _setext.match(lines[i+1] or ""): break
            buf.append(lines[i]); i+=1
        # Сохраняем ссылки в параграфах
        para_text = "\n".join(buf).strip()
        yield ("para", para_text)
        while i<n and lines[i].strip()=="": i+=1

def parse_markdown_to_tree(md: str) -> Dict[str,Any]:
    """node = {'level':int,'title':str,'content':str,'children':[node,...]}
    Сохраняет ссылки в формате [текст](url) в content и title"""
    root={"level":0,"title":"","content":"","children":[]}
    stack=[root]; pending=[]
    def attach():
        if pending:
            cur=stack[-1]
            cur["content"] += (("\n\n" if cur["content"] else "") + "\n\n".join(pending).strip())
            pending.clear()
    for blk in _iter_markdown_blocks(md):
        if blk[0]=="heading":
            attach()
            _, level, title = blk
            while stack and stack[-1]["level"]>=level: stack.pop()
            node={"level":level,"title":title,"content":"","children":[]}
            stack[-1]["children"].append(node); stack.append(node)
        else:
            _, txt = blk
            if txt: pending.append(txt)
    attach()
    
    # Постобработка: объединяем узлы с одинаковым уровнем, если они являются последовательными параграфами
    tree = _merge_sequential_paragraphs(root)
    
    # Постобработка: разбиваем content корневого узла на статьи, если они там есть
    tree = _split_content_into_articles(tree)
    
    # Постобработка: объединяем пункты под заголовками статей/разделов
    tree = _attach_items_to_headings(tree)
    
    # Постобработка: разбиваем подпункты в content на отдельные узлы
    tree = _split_sub_items_in_content(tree)
    
    # Постобработка: удаляем пустые промежуточные узлы
    tree = _remove_empty_intermediate_nodes(tree)
    
    return tree

def _is_roman_numeral(text: str) -> bool:
    """Проверяет, является ли текст римской цифрой"""
    if not text:
        return False
    text_upper = text.upper().strip()
    # Римские цифры: I, II, III, IV, V, VI, VII, VIII, IX, X, и т.д.
    roman_pattern = r'^[IVXLCDM]+$'
    return bool(re.match(roman_pattern, text_upper))

def _is_real_heading(title: str) -> bool:
    """Определяет, является ли title настоящим заголовком (РАЗДЕЛ, ГЛАВА, СТАТЬЯ и т.д.)"""
    if not title:
        return False
    
    title_upper = title.upper().strip()
    # Паттерны для настоящих заголовков
    heading_patterns = [
        r'^(РАЗДЕЛ|ГЛАВА|ЧАСТЬ|СТАТЬЯ|ПАРАГРАФ)\s+',
        r'^ГЛАВА\s+\d+',
        r'^РАЗДЕЛ\s+(ПЕРВЫЙ|ВТОРОЙ|ТРЕТИЙ|ЧЕТВЕРТЫЙ|ПЯТЫЙ|ШЕСТОЙ|СЕДЬМОЙ|ВОСЬМОЙ|ДЕВЯТЫЙ|ДЕСЯТЫЙ)',
        r'^КОНСТИТУЦИЯ',
        r'^ПРЕАМБУЛА',
        r'^[IVXLCDM]+\.',  # Римские цифры с точкой: I., II., III., IV., V.
    ]
    
    for pattern in heading_patterns:
        if re.match(pattern, title_upper):
            return True
    
    # Проверяем, начинается ли с римской цифры и точки
    parts = title_upper.split('.', 1)
    if len(parts) >= 2 and _is_roman_numeral(parts[0]):
        return True
    
    return False

def _is_likely_paragraph(node: Dict[str,Any]) -> bool:
    """Определяет, является ли узел скорее параграфом, чем заголовком"""
    title = node.get("title", "").strip()
    content = node.get("content", "").strip()
    has_children = bool(node.get("children"))
    
    # Если есть дочерние элементы - это заголовок
    if has_children:
        return False
    
    # Если нет title - это параграф
    if not title:
        return True
    
    # Если это настоящий заголовок - не параграф
    if _is_real_heading(title):
        return False
    
    # Если title очень длинный (>150 символов) - скорее всего параграф
    if len(title) > 150:
        return True
    
    # Если title заканчивается запятой или точкой с запятой - скорее всего часть предложения
    if title.endswith(",") or title.endswith(";"):
        return True
    
    # Если title начинается с маленькой буквы (не заглавной) - скорее всего часть предложения
    if title and title[0].islower():
        return True
    
    # Если есть content и title короткий - это заголовок
    if content and len(title) < 100:
        return False
    
    # Если нет content и title короткий - возможно заголовок, но если очень короткий и заканчивается запятой - параграф
    if not content and len(title) < 50 and (title.endswith(",") or title.endswith(";")):
        return True
    
    return False

def _merge_sequential_paragraphs(node: Dict[str,Any]) -> Dict[str,Any]:
    """Объединяет последовательные узлы одного уровня, если они являются параграфами"""
    if not node.get("children"):
        return node
    
    # Рекурсивно обрабатываем дочерние элементы
    processed_children = [_merge_sequential_paragraphs(ch) for ch in node.get("children", [])]
    
    merged_children = []
    i = 0
    while i < len(processed_children):
        current = processed_children[i].copy()
        
        # Если текущий узел - параграф, пытаемся объединить с последующими параграфами того же уровня
        if _is_likely_paragraph(current):
            # Собираем все последующие параграфы того же уровня
            paragraphs_to_merge = [current]
            j = i + 1
            while j < len(processed_children):
                next_node = processed_children[j]
                if (_is_likely_paragraph(next_node) and 
                    next_node.get("level", 0) == current.get("level", 0)):
                    paragraphs_to_merge.append(next_node)
                    j += 1
                else:
                    break
            
            # Объединяем все параграфы в один
            if len(paragraphs_to_merge) > 1:
                merged = {
                    "level": current.get("level", 1),
                    "title": "",
                    "content": "",
                    "children": []
                }
                
                # Собираем title и content из всех параграфов
                parts = []
                for para in paragraphs_to_merge:
                    para_title = para.get("title", "").strip()
                    para_content = para.get("content", "").strip()
                    
                    if para_title:
                        parts.append(para_title)
                    if para_content:
                        parts.append(para_content)
                
                # Первый непустой title становится title узла, остальное - content
                for para in paragraphs_to_merge:
                    para_title = para.get("title", "").strip()
                    if para_title and not merged["title"]:
                        # Проверяем, не является ли это частью предложения
                        if not _is_likely_paragraph(para) or len(para_title) < 50:
                            merged["title"] = para_title
                            break
                
                # Объединяем все части в content
                # Если title заканчивается запятой - объединяем через пробел, иначе через \n\n
                if parts:
                    # Проверяем, заканчивается ли первый title запятой
                    first_part = parts[0] if parts else ""
                    if first_part.endswith(",") or first_part.endswith(";"):
                        # Объединяем через пробел для продолжения предложения
                        merged["content"] = " ".join(parts).strip()
                    else:
                        merged["content"] = "\n\n".join(parts).strip()
                
                # Если title не установлен, но есть короткий первый title - используем его
                if not merged["title"] and paragraphs_to_merge[0].get("title"):
                    first_title = paragraphs_to_merge[0].get("title", "").strip()
                    # Если title заканчивается запятой - это часть предложения, не используем как title
                    if len(first_title) < 100 and not (first_title.endswith(",") or first_title.endswith(";")):
                        merged["title"] = first_title
                        # Убираем его из content, если он там есть
                        if merged["content"].startswith(first_title):
                            merged["content"] = merged["content"][len(first_title):].strip()
                
                merged_children.append(merged)
                i = j  # Пропускаем все объединенные узлы
                continue
        
        merged_children.append(current)
        i += 1
    
    node["children"] = merged_children
    return node

def _is_numbered_item(title: str) -> bool:
    """Определяет, является ли title пронумерованным пунктом (начинается с цифры и точки)"""
    if not title:
        return False
    
    title_stripped = title.strip()
    # Паттерн: начинается с цифры(ов), затем точка, затем пробел и текст
    # Примеры: "1. Текст", "1.1. Текст", "1) Текст"
    numbered_patterns = [
        r'^\d+[\.\)]\s+',  # "1. " или "1) "
        r'^\d+\.\d+[\.\)]\s+',  # "1.1. " или "1.1) "
        r'^[а-яА-Я]\.\s+',  # "а. " (буквенная нумерация)
    ]
    
    for pattern in numbered_patterns:
        if re.match(pattern, title_stripped):
            return True
    
    return False

def _is_sub_item(title: str) -> bool:
    """Определяет, является ли title подпунктом (начинается с цифры и скобки, например "1)")"""
    if not title:
        return False
    
    title_stripped = title.strip()
    # Подпункты обычно имеют формат "1)", "2)", "а)", "б)" и т.д.
    sub_item_patterns = [
        r'^\d+\)\s+',  # "1) "
        r'^[а-яА-Я]\)\s+',  # "а) " (буквенная нумерация)
    ]
    
    for pattern in sub_item_patterns:
        if re.match(pattern, title_stripped):
            return True
    
    return False

def _extract_sub_items_from_content(content: str) -> List[Dict[str,Any]]:
    """Извлекает подпункты вида "1)", "2)", "3)" из content, если они идут в одну строку через точку с запятой"""
    if not content:
        return []
    
    items = []
    
    # Ищем подпункты вида "1) текст; 2) текст; 3) текст;"
    # Паттерн: цифра + скобка + пробел + текст до следующего подпункта или конца строки
    # Учитываем, что текст может содержать точки с запятой внутри (но не перед следующим подпунктом)
    # Более точный паттерн: ищем "цифра) текст" до следующего "цифра)" или конца строки
    sub_item_pattern = re.compile(r'(\d+\)\s+[^;]*?)(?=\s*;\s*\d+\)|;?\s*$)', re.MULTILINE | re.DOTALL)
    
    matches = list(sub_item_pattern.finditer(content))
    
    # Если нашли несколько подпунктов
    if len(matches) > 1:
        # Проверяем, что они действительно в одной строке (нет \n между ними, или только один)
        first_match = matches[0]
        last_match = matches[-1]
        text_between = content[first_match.start():last_match.end()]
        
        # Если между первым и последним подпунктом нет переносов строк (или только один), значит они в одной строке
        if '\n' not in text_between or text_between.count('\n') <= 1:
            # Разбиваем на отдельные подпункты
            for match in matches:
                sub_item_text = match.group(1).strip()
                # Удаляем точку с запятой в конце, если есть
                if sub_item_text.endswith(';'):
                    sub_item_text = sub_item_text[:-1].strip()
                
                if sub_item_text:
                    items.append({
                        "level": 4,  # Будет установлен правильно позже
                        "title": sub_item_text,
                        "content": "",
                        "children": []
                    })
    elif len(matches) == 1:
        # Если нашли только один подпункт, но в content есть другие подпункты через точку с запятой
        # Проверяем, есть ли в content другие подпункты после точки с запятой
        first_match = matches[0]
        text_after = content[first_match.end():].strip()
        # Ищем другие подпункты после первого
        if text_after and re.search(r';\s*\d+\)', text_after):
            # Есть другие подпункты, извлекаем все через более простой метод
            # Разбиваем по паттерну "цифра) текст;"
            all_matches = re.finditer(r'(\d+\)\s+[^;]+?)(?=\s*;\s*\d+\)|;?\s*$)', content, re.MULTILINE | re.DOTALL)
            for match in all_matches:
                sub_item_text = match.group(1).strip()
                if sub_item_text.endswith(';'):
                    sub_item_text = sub_item_text[:-1].strip()
                if sub_item_text:
                    items.append({
                        "level": 4,
                        "title": sub_item_text,
                        "content": "",
                        "children": []
                    })
    
    return items

def _is_letter_sub_item(title: str) -> bool:
    """Определяет, является ли title буквенным подпунктом (начинается с буквы и скобки, например "а)")"""
    if not title:
        return False
    
    title_stripped = title.strip()
    # Буквенные подпункты: "а)", "б)", "в)" и т.д.
    letter_pattern = re.compile(r'^[а-яА-Я]\)\s+')
    return bool(letter_pattern.match(title_stripped))

def _extract_letter_sub_items_from_content(content: str) -> List[Dict[str,Any]]:
    """Извлекает буквенные подпункты (а), б), в)) из content текста"""
    if not content:
        return []
    
    items = []
    
    # Находим все позиции, где начинается новый буквенный подпункт
    # Паттерн: начало строки, \n\n, или \n, затем буква, скобка, пробел
    letter_sub_item_pattern = re.compile(r'(?:^|\n\n|\n)([а-яА-Я]\)\s+)', re.MULTILINE)
    
    # Находим все начала подпунктов
    start_positions = []
    for match in letter_sub_item_pattern.finditer(content):
        pos = match.start()
        match_text = match.group(0)
        # Пропускаем разделители (\n\n или \n)
        if match_text.startswith('\n\n'):
            pos = match.start() + 2
        elif match_text.startswith('\n'):
            pos = match.start() + 1
        start_positions.append(pos)
    
    # Если нашли начала подпунктов, разбиваем content по ним
    if start_positions:
        for i, start_pos in enumerate(start_positions):
            # Определяем конец текущего подпункта (начало следующего или конец content)
            end_pos = start_positions[i + 1] if i + 1 < len(start_positions) else len(content)
            
            # Извлекаем текст подпункта
            item_text = content[start_pos:end_pos].strip()
            if not item_text:
                continue
            
            # Разделяем на первую строку (title) и остальное (content)
            lines = item_text.split("\n")
            first_line = lines[0].strip() if lines else ""
            
            if _is_letter_sub_item(first_line):
                # Остальной текст - это content подпункта
                remaining_lines = lines[1:] if len(lines) > 1 else []
                item_content = "\n".join(remaining_lines).strip()
                
                items.append({
                    "level": 2,  # Будет установлен правильно позже
                    "title": first_line,
                    "content": item_content,
                    "children": []
                })
    
    return items

def _is_main_item(title: str) -> bool:
    """Определяет, является ли title основным пунктом (начинается с цифры и точки, например "1.")"""
    if not title:
        return False
    
    title_stripped = title.strip()
    # Основные пункты обычно имеют формат "1.", "2.", "а." и т.д.
    main_item_patterns = [
        r'^\d+\.\s+',  # "1. "
        r'^[а-яА-Я]\.\s+',  # "а. " (буквенная нумерация)
    ]
    
    for pattern in main_item_patterns:
        if re.match(pattern, title_stripped):
            return True
    
    return False

def _is_article_heading(title: str) -> bool:
    """Определяет, является ли title заголовком статьи/раздела/главы, под которым могут быть пункты"""
    if not title:
        return False
    
    title_stripped = title.strip()
    title_upper = title_stripped.upper()
    
    # Паттерны для заголовков, под которыми обычно идут пункты
    # Учитываем варианты с пробелом и без, с точкой и без
    article_patterns = [
        r'^(СТАТЬЯ|ARTICLE)\s*\d+',  # "Статья 1" или "Статья1"
        r'^(РАЗДЕЛ|SECTION)\s*\d+',
        r'^(ГЛАВА|CHAPTER)\s*\d+',
        r'^(ЧАСТЬ|PART)\s*\d+',
        r'^(ПАРАГРАФ|PARAGRAPH)\s*\d+',
        r'^СТАТЬЯ\s*\d+',
        r'^ГЛАВА\s*\d+',
        r'^РАЗДЕЛ\s*\d+',
        r'^[IVXLCDM]+\.',  # Римские цифры с точкой: "I.", "II.", "III.", "IV.", "V."
    ]
    
    for pattern in article_patterns:
        if re.match(pattern, title_upper):
            return True
    
    # Проверяем, начинается ли с римской цифры и точки
    parts = title_upper.split('.', 1)
    if len(parts) >= 2 and _is_roman_numeral(parts[0]):
        return True
    
    return False

def _extract_numbered_items_from_content_with_positions(content: str) -> Tuple[List[Dict[str,Any]], List[Tuple[int, int]]]:
    """Извлекает пронумерованные пункты из content текста и возвращает их позиции"""
    if not content:
        return [], []
    
    items = []
    positions = []
    
    # Улучшенный паттерн для поиска пунктов:
    # 1. Пункты могут начинаться с начала строки, после \n\n, после \n, или после пробела/точки с запятой
    # 2. Формат: цифра + скобка/точка + пробел + текст
    # 3. Пункт заканчивается перед следующим пунктом или в конце строки/текста
    
    # Находим все позиции, где начинается новый пронумерованный пункт
    # Паттерн: начало строки, \n\n, \n, пробел, или точка с запятой + пробел, затем цифра, скобка, пробел
    # Также учитываем случаи, когда пункты идут подряд через точку с запятой: "1) текст; 2) текст;"
    numbered_start_pattern = re.compile(r'(?:^|\n\n|\n|;\s*|\.\s+)(\d+\))\s+', re.MULTILINE)
    
    # Находим все начала пунктов
    start_positions = []
    for match in numbered_start_pattern.finditer(content):
        # Позиция начала пункта (после разделителя)
        pos = match.start()
        match_text = match.group(0)
        # Определяем, сколько символов пропустить
        if match_text.startswith('\n\n'):
            pos = match.start() + 2
        elif match_text.startswith('\n'):
            pos = match.start() + 1
        elif match_text.startswith('; '):
            pos = match.start() + 2
        elif match_text.startswith(';'):
            pos = match.start() + 1
        elif match_text.startswith('. '):
            pos = match.start() + 2
        # Если начинается с начала строки, pos уже правильный
        start_positions.append((pos, match.end()))
    
    # Если нашли начала пунктов, разбиваем content по ним
    if start_positions:
        for i, (start_pos, match_end) in enumerate(start_positions):
            # Определяем конец текущего пункта (начало следующего или конец content)
            if i + 1 < len(start_positions):
                next_start_pos, _ = start_positions[i + 1]
                end_pos = next_start_pos
            else:
                end_pos = len(content)
            
            # Извлекаем текст пункта
            item_text = content[start_pos:end_pos].strip()
            if not item_text:
                continue
            
            # Определяем конец текущего пункта
            # Пункт заканчивается перед следующим пунктом (который начинается с "цифра)")
            # Или перед точкой с запятой, если следующий пункт идет сразу после
            item_end_pos = len(item_text)
            
            # Ищем следующий пункт в тексте (паттерн: пробел или точка с запятой + пробел + цифра + скобка)
            next_item_match = re.search(r';\s*(\d+\))\s+', item_text)
            if next_item_match:
                # Найден следующий пункт, текущий заканчивается перед ним
                item_end_pos = next_item_match.start()
                item_text = item_text[:item_end_pos].strip()
            
            # Удаляем точку с запятой в конце, если есть
            if item_text.endswith(';'):
                item_text = item_text[:-1].strip()
            
            if not item_text:
                continue
            
            # Разделяем на первую строку (title) и остальное (content)
            # Пункт может быть на одной строке или на нескольких
            lines = item_text.split("\n")
            first_line = lines[0].strip() if lines else ""
            
            if not first_line:
                continue
            
            # Проверяем, является ли первая строка пронумерованным пунктом
            if _is_numbered_item(first_line):
                # Остальной текст - это content пункта
                remaining_lines = lines[1:] if len(lines) > 1 else []
                item_content = "\n".join(remaining_lines).strip()
                
                # Вычисляем реальные позиции в исходном content
                item_start = start_pos
                # Конец пункта - это начало следующего пункта или конец текста
                if next_item_match:
                    item_end = start_pos + item_end_pos + 1  # +1 чтобы включить точку с запятой
                else:
                    item_end = end_pos
                
                items.append({
                    "level": 2,  # Будет установлен правильно позже
                    "title": first_line,
                    "content": item_content,
                    "children": []
                })
                positions.append((item_start, item_end))
    
    return items, positions

def _extract_numbered_items_from_content(content: str) -> List[Dict[str,Any]]:
    """Извлекает пронумерованные пункты из content текста"""
    items, _ = _extract_numbered_items_from_content_with_positions(content)
    if items:
        return items
    
    # Fallback логика, если основная функция не нашла пункты
    items = []
    
    # Улучшенный паттерн для поиска пунктов:
    # 1. Пункты могут начинаться с начала строки, после \n\n, после \n, или после пробела/точки с запятой
    # 2. Формат: цифра + скобка/точка + пробел + текст
    # 3. Пункт заканчивается перед следующим пунктом или в конце строки/текста
    
    # Находим все позиции, где начинается новый пронумерованный пункт
    # Паттерн: начало строки, \n\n, \n, пробел, или точка с запятой + пробел, затем цифра, скобка, пробел
    # Также учитываем случаи, когда пункты идут подряд через точку с запятой: "1) текст; 2) текст;"
    numbered_start_pattern = re.compile(r'(?:^|\n\n|\n|;\s*|\.\s+)(\d+\))\s+', re.MULTILINE)
    
    # Находим все начала пунктов
    start_positions = []
    for match in numbered_start_pattern.finditer(content):
        # Позиция начала пункта (после разделителя)
        pos = match.start()
        match_text = match.group(0)
        # Определяем, сколько символов пропустить
        if match_text.startswith('\n\n'):
            pos = match.start() + 2
        elif match_text.startswith('\n'):
            pos = match.start() + 1
        elif match_text.startswith('; '):
            pos = match.start() + 2
        elif match_text.startswith('. '):
            pos = match.start() + 2
        # Если начинается с начала строки, pos уже правильный
        start_positions.append((pos, match.end()))
    
    # Если нашли начала пунктов, разбиваем content по ним
    if start_positions:
        for i, (start_pos, match_end) in enumerate(start_positions):
            # Определяем конец текущего пункта (начало следующего или конец content)
            if i + 1 < len(start_positions):
                next_start_pos, _ = start_positions[i + 1]
                end_pos = next_start_pos
            else:
                end_pos = len(content)
            
            # Извлекаем текст пункта
            item_text = content[start_pos:end_pos].strip()
            if not item_text:
                continue
            
            # Определяем конец текущего пункта
            # Пункт заканчивается перед следующим пунктом (который начинается с "цифра)")
            # Или перед точкой с запятой, если следующий пункт идет сразу после
            item_end_pos = len(item_text)
            
            # Ищем следующий пункт в тексте (паттерн: пробел или точка с запятой + пробел + цифра + скобка)
            next_item_match = re.search(r';\s*(\d+\))\s+', item_text)
            if next_item_match:
                # Найден следующий пункт, текущий заканчивается перед ним
                item_end_pos = next_item_match.start()
                item_text = item_text[:item_end_pos].strip()
            
            # Удаляем точку с запятой в конце, если есть
            if item_text.endswith(';'):
                item_text = item_text[:-1].strip()
            
            if not item_text:
                continue
            
            # Разделяем на первую строку (title) и остальное (content)
            # Пункт может быть на одной строке или на нескольких
            lines = item_text.split("\n")
            first_line = lines[0].strip() if lines else ""
            
            if not first_line:
                continue
            
            # Проверяем, является ли первая строка пронумерованным пунктом
            if _is_numbered_item(first_line):
                # Остальной текст - это content пункта
                remaining_lines = lines[1:] if len(lines) > 1 else []
                item_content = "\n".join(remaining_lines).strip()
                
                items.append({
                    "level": 2,  # Будет установлен правильно позже
                    "title": first_line,
                    "content": item_content,
                    "children": []
                })
    else:
        # Если не нашли через паттерн, пробуем разделить по точкам с запятой
        # (для случаев, когда пункты в одной строке: "1) текст; 2) текст;")
        if ';' in content and re.search(r'\d+\)\s+', content):
            # Ищем все пункты в формате "цифра) текст;"
            item_matches = list(re.finditer(r'(\d+\))\s+([^;]+?)(?=;\s*\d+\)|;?\s*$)', content))
            
            if item_matches:
                for match in item_matches:
                    item_num = match.group(1)
                    item_text = match.group(2).strip()
                    
                    # Удаляем точку с запятой в конце, если есть
                    if item_text.endswith(';'):
                        item_text = item_text[:-1].strip()
                    
                    if item_text:
                        item_title = f"{item_num} {item_text}"
                        
                        if _is_numbered_item(item_title):
                            items.append({
                                "level": 2,
                                "title": item_title,
                                "content": "",
                                "children": []
                            })
        
        # Если не нашли через точки с запятой, пробуем разделить по \n\n
        if not items:
            parts = content.split("\n\n")
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                
                lines = part.split("\n")
                first_line = lines[0].strip() if lines else ""
                
                if _is_numbered_item(first_line):
                    remaining_lines = lines[1:] if len(lines) > 1 else []
                    item_content = "\n".join(remaining_lines).strip()
                    
                    items.append({
                        "level": 2,
                        "title": first_line,
                        "content": item_content,
                        "children": []
                    })
    
    return items

def _split_content_into_articles(node: Dict[str,Any]) -> Dict[str,Any]:
    """Разбивает content корневого узла на статьи, если они там есть"""
    # Обрабатываем только корневой узел (level 0) с пустым title
    if node.get("level", 0) != 0 or node.get("title", "").strip():
        # Рекурсивно обрабатываем дочерние элементы
        if node.get("children"):
            node["children"] = [_split_content_into_articles(ch) for ch in node["children"]]
        return node
    
    content = node.get("content", "").strip()
    if not content:
        # Рекурсивно обрабатываем дочерние элементы
        if node.get("children"):
            node["children"] = [_split_content_into_articles(ch) for ch in node["children"]]
        return node
    
    # Ищем паттерны "Статья X" в content
    # Паттерн: начало строки или два переноса строки, затем "Статья" + пробел + цифра, затем один или два переноса строки
    # Учитываем разные варианты: "\n\nСтатья 1\n\n", "\nСтатья 1\n", "^Статья 1\n" (в начале)
    article_pattern = re.compile(r'(?:^|\n\n|\n)(Статья\s+\d+)\s*(?:\n\n|\n|$)', re.MULTILINE | re.IGNORECASE)
    
    # Находим все совпадения
    matches = list(article_pattern.finditer(content))
    
    if not matches:
        # Если не нашли статьи в content, рекурсивно обрабатываем дочерние элементы
        if node.get("children"):
            node["children"] = [_split_content_into_articles(ch) for ch in node["children"]]
        return node
    
    # Разбиваем content на части
    extracted_articles = []
    remaining_content_parts = []
    
    # Текст до первой статьи
    first_match = matches[0]
    text_before = content[:first_match.start()].strip()
    if text_before:
        remaining_content_parts.append(text_before)
    
    # Обрабатываем каждую статью
    for i, match in enumerate(matches):
        article_title = match.group(1)  # "Статья X"
        
        # Определяем начало content статьи (после "Статья X" и переносов строк)
        article_start = match.end()
        # Пропускаем пробелы и переносы строк после "Статья X"
        while article_start < len(content) and content[article_start] in ['\n', ' ', '\t']:
            article_start += 1
        
        # Конец статьи - это начало следующей статьи или конец content
        if i + 1 < len(matches):
            next_match = matches[i + 1]
            article_end = next_match.start()
        else:
            article_end = len(content)
        
        article_content = content[article_start:article_end].strip()
        
        # Создаем узел для статьи
        article_node = {
            "level": 2,  # Уровень статьи
            "title": article_title,
            "content": article_content,
            "children": []
        }
        extracted_articles.append(article_node)
    
    # Обновляем корневой узел
    if remaining_content_parts:
        node["content"] = "\n\n".join(remaining_content_parts)
    else:
        node["content"] = ""
    
    # Вставляем извлеченные статьи в начало children
    existing_children = node.get("children", [])
    node["children"] = extracted_articles + existing_children
    
    # Рекурсивно обрабатываем все дочерние элементы (включая только что созданные статьи)
    if node.get("children"):
        node["children"] = [_split_content_into_articles(ch) for ch in node["children"]]
    
    return node

def _split_sub_items_in_content(node: Dict[str,Any]) -> Dict[str,Any]:
    """Разбивает подпункты вида "1)", "2)", "3)" в content на отдельные узлы, если они идут в одну строку"""
    # Рекурсивно обрабатываем дочерние элементы сначала
    if node.get("children"):
        node["children"] = [_split_sub_items_in_content(ch) for ch in node["children"]]
    
    # Проверяем content текущего узла на наличие подпунктов в одну строку
    content = node.get("content", "").strip()
    title = node.get("title", "").strip()
    
    # Если title начинается с подпункта вида "1)", проверяем, есть ли другие подпункты в content
    if _is_sub_item(title) and content:
        # Проверяем, есть ли в content другие подпункты
        # Объединяем title и content для проверки (без добавления "; ", так как они могут быть разделены по-разному)
        # Если content начинается с подпункта, объединяем через "; ", иначе просто конкатенируем
        if _is_sub_item(content.strip().split(';')[0] if ';' in content else content.strip()):
            combined_text = title + "; " + content
        else:
            combined_text = title + " " + content if title else content
        sub_items = _extract_sub_items_from_content(combined_text)
        
        if sub_items and len(sub_items) > 1:
            # Если нашли несколько подпунктов, разбиваем их все на отдельные узлы
            # Удаляем title, так как он будет частью подпунктов
            node["title"] = ""
            
            # Удаляем все подпункты из combined_text
            # Создаем паттерн для удаления всех подпунктов сразу
            cleaned_content = combined_text
            # Удаляем каждый подпункт по отдельности, начиная с конца
            for item in reversed(sub_items):
                sub_item_text = item["title"]
                # Удаляем подпункт с точкой с запятой или без, учитывая возможные пробелы
                # Используем более точный паттерн, который учитывает, что подпункт может быть в начале строки
                pattern = r'\s*' + re.escape(sub_item_text) + r'\s*;?\s*'
                cleaned_content = re.sub(pattern, '', cleaned_content, count=1)
            
            # Очищаем множественные пробелы, точки с запятой и переносы строк
            cleaned_content = re.sub(r'\s*;\s*', ' ', cleaned_content)
            cleaned_content = re.sub(r'\s+', ' ', cleaned_content).strip()
            
            # Удаляем точки с запятой в начале и конце
            cleaned_content = cleaned_content.strip(';').strip()
            
            # Если остался только один подпункт (который не был удален), удаляем его тоже
            if _is_sub_item(cleaned_content.strip()):
                cleaned_content = ""
            
            if not cleaned_content or cleaned_content.isspace():
                node["content"] = ""
            else:
                node["content"] = cleaned_content
            
            # Устанавливаем правильный уровень для подпунктов
            current_level = node.get("level", 0)
            for item in sub_items:
                item["level"] = current_level + 1
            
            # Добавляем подпункты в children
            if not node.get("children"):
                node["children"] = []
            node["children"] = sub_items + node["children"]
            return node
    
    # Если title не является подпунктом, проверяем только content
    if not content:
        return node
    
    # Ищем подпункты вида "1) текст; 2) текст; 3) текст;"
    sub_items = _extract_sub_items_from_content(content)
    
    if sub_items and len(sub_items) > 1:
        # Если нашли несколько подпунктов в одну строку, разбиваем их
        sub_item_texts = [item["title"] for item in sub_items]
        
        # Удаляем подпункты из content
        cleaned_content = content
        for sub_item_text in reversed(sub_item_texts):
            pattern = re.escape(sub_item_text) + r'\s*;?\s*'
            cleaned_content = re.sub(pattern, '', cleaned_content, count=1)
        
        cleaned_content = re.sub(r'\s+', ' ', cleaned_content).strip()
        if not cleaned_content or cleaned_content.isspace():
            node["content"] = ""
        else:
            node["content"] = cleaned_content
        
        # Устанавливаем правильный уровень для подпунктов
        current_level = node.get("level", 0)
        for item in sub_items:
            item["level"] = current_level + 1
        
        # Добавляем подпункты в children
        if not node.get("children"):
            node["children"] = []
        node["children"] = sub_items + node["children"]
    
    return node

def _remove_empty_intermediate_nodes(node: Dict[str,Any]) -> Dict[str,Any]:
    """Удаляет пустые промежуточные узлы (с пустым title и content, но с children)"""
    # Рекурсивно обрабатываем дочерние элементы сначала
    if node.get("children"):
        processed_children = []
        for child in node["children"]:
            processed_child = _remove_empty_intermediate_nodes(child)
            # Если дочерний узел пустой (нет title и content), но есть children, поднимаем children на уровень выше
            if (not processed_child.get("title", "").strip() and 
                not processed_child.get("content", "").strip() and 
                processed_child.get("children")):
                # Поднимаем children на уровень выше
                # Устанавливаем level children равным level родителя + 1
                parent_level = node.get("level", 0)
                for grandchild in processed_child.get("children", []):
                    # Устанавливаем level как parent_level + 1 (дети должны быть на один уровень ниже родителя)
                    grandchild["level"] = parent_level + 1
                processed_children.extend(processed_child.get("children", []))
            else:
                processed_children.append(processed_child)
        node["children"] = processed_children
    
    return node

def _attach_items_to_headings(node: Dict[str,Any]) -> Dict[str,Any]:
    """Прикрепляет пронумерованные пункты к заголовкам статей/разделов"""
    if not node.get("children"):
        return node
    
    # Рекурсивно обрабатываем дочерние элементы
    processed_children = [_attach_items_to_headings(ch) for ch in node.get("children", [])]
    
    result_children = []
    i = 0
    
    while i < len(processed_children):
        current = processed_children[i].copy()
        current_title = current.get("title", "").strip()
        
        # Проверяем, является ли текущий узел заголовком статьи/раздела
        if _is_article_heading(current_title) and not current.get("children"):
            # Собираем все последующие элементы
            items_to_attach = []
            content_parts = []
            current_level = current.get("level", 0)
            
            # Сначала проверяем content самого заголовка на наличие пронумерованных пунктов
            current_content = current.get("content", "").strip()
            if current_content:
                # Проверяем, есть ли в content пронумерованные пункты (не только в первой строке)
                # Ищем паттерн пронумерованного пункта в content
                has_numbered_items = bool(re.search(r'\n\n\d+[\.\)]\s+', current_content) or 
                                         re.match(r'^\d+[\.\)]\s+', current_content))
                
                if has_numbered_items:
                    # Извлекаем пункты из content заголовка
                    # Используем улучшенную функцию, которая также возвращает позиции
                    result = _extract_numbered_items_from_content_with_positions(current_content)
                    # Защита от неправильной распаковки
                    if isinstance(result, tuple) and len(result) == 2:
                        extracted_items, item_positions = result
                    else:
                        # Если функция вернула неожиданное количество значений, используем только первый элемент
                        extracted_items = result[0] if isinstance(result, tuple) and len(result) > 0 else []
                        item_positions = result[1] if isinstance(result, tuple) and len(result) > 1 else []
                    if extracted_items and item_positions:
                        # Удаляем все извлеченные пункты из content используя позиции
                        cleaned_content = current_content
                        
                        # Удаляем пункты в обратном порядке, чтобы не сбить позиции
                        for item, (start_pos, end_pos) in zip(reversed(extracted_items), reversed(item_positions)):
                            # Удаляем пункт из content
                            before = cleaned_content[:start_pos].rstrip()
                            after = cleaned_content[end_pos:].lstrip()
                            
                            # Объединяем части, убирая лишние пробелы
                            if before and after:
                                cleaned_content = before + " " + after
                            elif before:
                                cleaned_content = before
                            elif after:
                                cleaned_content = after
                            else:
                                cleaned_content = ""
                        
                        # Очищаем множественные пробелы и пустые строки
                        cleaned_content = re.sub(r'\s+', ' ', cleaned_content)
                        cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)
                        cleaned_content = cleaned_content.strip()
                        
                        current["content"] = cleaned_content
                        
                        # Добавляем извлеченные пункты
                        for item in extracted_items:
                            item["level"] = current_level + 1
                            items_to_attach.append(item)
                    elif extracted_items:
                        # Если позиции не найдены, используем старую логику
                        # Удаляем пункты по тексту
                        cleaned_content = current_content
                        for item in reversed(extracted_items):
                            item_title = item.get("title", "").strip()
                            # Ищем и удаляем пункт
                            item_pos = cleaned_content.find(item_title)
                            if item_pos != -1:
                                end_pos = item_pos + len(item_title)
                                # Пропускаем точку с запятую и пробелы
                                while end_pos < len(cleaned_content) and cleaned_content[end_pos] in [';', ' ', '\n']:
                                    end_pos += 1
                                cleaned_content = cleaned_content[:item_pos].rstrip() + " " + cleaned_content[end_pos:].lstrip()
                        
                        cleaned_content = re.sub(r'\s+', ' ', cleaned_content).strip()
                        current["content"] = cleaned_content
                        
                        for item in extracted_items:
                            item["level"] = current_level + 1
                            items_to_attach.append(item)
            
            j = i + 1
            last_item_number = None  # Отслеживаем номер последнего пункта
            
            while j < len(processed_children):
                next_node = processed_children[j]
                next_title = next_node.get("title", "").strip()
                next_content = next_node.get("content", "").strip()
                next_level = next_node.get("level", 0)
                
                # ВАЖНО: Проверяем, является ли следующий узел заголовком статьи/раздела
                # Это должно быть ПЕРВОЙ проверкой, чтобы не пропустить новую статью
                # Проверяем на том же уровне или выше - это точно новая статья/раздел
                if _is_article_heading(next_title):
                    # Если это заголовок статьи на том же или более высоком уровне - прекращаем
                    if next_level <= current_level:
                        break
                    # Если это заголовок статьи на более низком уровне, но мы уже собрали пункты - тоже прекращаем
                    # (может быть подраздел внутри статьи)
                    if next_level > current_level and items_to_attach:
                        break
                
                # Если следующий узел - пронумерованный пункт в title
                if _is_numbered_item(next_title):
                    # Проверяем, не начинается ли пункт с "1." - это может означать новую статью
                    if next_title.strip().startswith("1.") and last_item_number is not None and last_item_number > 1:
                        # Если мы уже собрали пункты с номерами > 1, а следующий пункт начинается с "1.",
                        # это может быть начало новой статьи - прекращаем
                        break
                    
                    if next_level >= current_level and next_level <= current_level + 1:
                        # Извлекаем номер пункта для отслеживания
                        num_match = re.match(r'^(\d+)[\.\)]', next_title)
                        if num_match:
                            item_num = int(num_match.group(1))
                            # Если пункт начинается с "1." и мы уже собрали пункты, это новая статья
                            if item_num == 1 and last_item_number is not None and last_item_number > 1:
                                break
                            last_item_number = item_num
                        
                        item_copy = next_node.copy()
                        item_copy["level"] = current_level + 1
                        items_to_attach.append(item_copy)
                        j += 1
                        continue
                    else:
                        break
                
                # Если следующий узел имеет content с пронумерованными пунктами
                # Проверяем, есть ли в content пронумерованные пункты
                if next_content:
                    # Проверяем, начинается ли content с пронумерованного пункта
                    content_lines = next_content.split("\n")
                    first_content_line = content_lines[0].strip() if content_lines else ""
                    if _is_numbered_item(first_content_line):
                        # Извлекаем пункты из content
                        extracted_items = _extract_numbered_items_from_content(next_content)
                        if extracted_items:
                            # Проверяем первый пункт - если он начинается с "1." и мы уже собрали пункты, это новая статья
                            first_item_title = extracted_items[0].get("title", "").strip()
                            if first_item_title.startswith("1.") and last_item_number is not None and last_item_number > 1:
                                break
                            
                            for item in extracted_items:
                                item["level"] = current_level + 1
                                items_to_attach.append(item)
                                # Обновляем last_item_number
                                num_match = re.match(r'^(\d+)[\.\)]', item.get("title", ""))
                                if num_match:
                                    last_item_number = int(num_match.group(1))
                            j += 1
                            # Продолжаем собирать следующие пункты
                            continue
                
                # Если следующий узел не является пронумерованным пунктом
                # Проверяем различные случаи:
                # 1. Уровень меньше текущего - это возврат к более высокому уровню, прекращаем
                if next_level < current_level:
                    break
                # 2. Это обычный текст на том же уровне - добавляем в content статьи
                if next_level == current_level:
                    # Если это не пронумерованный пункт, добавляем в content
                    if not _is_numbered_item(next_title):
                        text_to_add = next_title if next_title else next_content
                        if text_to_add:
                            if current.get("content"):
                                current["content"] += "\n\n" + text_to_add
                            else:
                                current["content"] = text_to_add
                        j += 1
                        continue
                    else:
                        break
                # 3. Если уровень больше, но это не пункт - пропускаем (может быть подпункт)
                j += 1
            
            # Если нашли пункты для прикрепления
            if items_to_attach:
                # Обрабатываем вложенные пункты (подпункты) - прикрепляем подпункты (1), 2), 3)) к основным пунктам (1., 2., 3.)
                processed_items = []
                used_sub_items = set()  # Отслеживаем подпункты, которые уже были использованы
                
                for item_idx, item in enumerate(items_to_attach):
                    item_copy = item.copy()
                    item_title = item.get("title", "").strip()
                    item_level = item.get("level", 0)
                    
                    # Пропускаем подпункты, которые уже были прикреплены к предыдущим основным пунктам
                    if _is_sub_item(item_title) and id(item) in used_sub_items:
                        continue
                    
                    # Проверяем content пункта на наличие буквенных подпунктов (а), б), в))
                    item_content = item_copy.get("content", "").strip()
                    if item_content:
                        letter_sub_items = _extract_letter_sub_items_from_content(item_content)
                        if letter_sub_items:
                            # Разделяем content на текст до подпунктов и сами подпункты
                            first_sub_item_title = letter_sub_items[0].get("title", "").strip()
                            first_sub_item_pos = item_content.find(first_sub_item_title)
                            
                            if first_sub_item_pos > 0:
                                # Сохраняем текст до первого подпункта в content
                                text_before_sub_items = item_content[:first_sub_item_pos].strip()
                                item_copy["content"] = text_before_sub_items
                            else:
                                # Если подпункты начинаются с самого начала, очищаем content
                                item_copy["content"] = ""
                            
                            # Устанавливаем правильный уровень для буквенных подпунктов
                            for sub_item in letter_sub_items:
                                sub_item["level"] = item_level + 1
                            
                            # Добавляем буквенные подпункты как children
                            if not item_copy.get("children"):
                                item_copy["children"] = []
                            item_copy["children"].extend(letter_sub_items)
                    
                    # Проверяем, является ли это основным пунктом (1., 2., 3.)
                    if _is_main_item(item_title):
                        # Ищем подпункты после этого основного пункта
                        # Сначала ищем среди следующих пунктов в items_to_attach
                        sub_items = []
                        for k in range(item_idx + 1, len(items_to_attach)):
                            next_item = items_to_attach[k]
                            next_title = next_item.get("title", "").strip()
                            
                            # Если это подпункт (1), 2), 3))
                            if _is_sub_item(next_title):
                                sub_item_copy = next_item.copy()
                                sub_item_copy["level"] = item_level + 1
                                sub_items.append(sub_item_copy)
                                used_sub_items.add(id(next_item))
                            # Если это следующий основной пункт - прекращаем
                            elif _is_main_item(next_title):
                                break
                            # Если это не подпункт и не основной пункт - прекращаем
                            else:
                                break
                        
                        # Также ищем подпункты в следующих узлах processed_children, которые еще не были обработаны
                        for k in range(j, len(processed_children)):
                            next_node = processed_children[k]
                            next_title = next_node.get("title", "").strip()
                            next_level = next_node.get("level", 0)
                            
                            # Если это подпункт (1), 2), 3))
                            if _is_sub_item(next_title):
                                # Проверяем, что уровень подходит
                                if next_level >= item_level and next_level <= item_level + 1:
                                    sub_item_copy = next_node.copy()
                                    sub_item_copy["level"] = item_level + 1
                                    sub_items.append(sub_item_copy)
                                    continue
                                else:
                                    break
                            
                            # Если это следующий основной пункт - прекращаем
                            if _is_main_item(next_title):
                                break
                            
                            # Если это заголовок статьи - прекращаем
                            if _is_article_heading(next_title):
                                break
                            
                            # Если уровень меньше - это возврат к более высокому уровню
                            if next_level < item_level:
                                break
                            
                            # Если это не подпункт и не основной пункт, прекращаем поиск подпунктов
                            if not _is_numbered_item(next_title):
                                break
                        
                        # Если нашли подпункты, добавляем их как children
                        if sub_items:
                            item_copy["children"] = sub_items
                        else:
                            item_copy["children"] = []
                    else:
                        item_copy["children"] = []
                    
                    processed_items.append(item_copy)
                
                current["children"] = processed_items
                result_children.append(current)
                i = j  # Пропускаем все прикрепленные узлы
                continue
            # Если добавили content, но не было пунктов
            elif current.get("content"):
                result_children.append(current)
                i = j if j > i + 1 else i + 1
                continue
        
        result_children.append(current)
        i += 1
    
    node["children"] = result_children
    return node

# ---------- 2) Система очистки текста ----------

def load_cleaning_patterns(pattern_file: str = "возможные_паттерны_для_очистки.json") -> List[Dict[str, str]]:
    """Загружает паттерны для очистки из JSON файла"""
    import os
    import json
    
    # Пытаемся найти файл в текущей директории или рядом со скриптом
    script_dir = os.path.dirname(os.path.abspath(__file__))
    pattern_path = os.path.join(script_dir, pattern_file)
    
    if not os.path.exists(pattern_path):
        # Если файл не найден, возвращаем пустой список
        print(f"[WARNING] Файл с паттернами очистки не найден: {pattern_path}")
        return []
    
    try:
        with open(pattern_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data.get("patterns", [])
    except Exception as e:
        print(f"[WARNING] Ошибка при загрузке паттернов очистки: {e}")
        return []

def clean_text_with_patterns(text: str, patterns: List[Dict[str, str]] = None) -> str:
    """Очищает текст используя паттерны из JSON файла"""
    if patterns is None:
        patterns = load_cleaning_patterns()
    
    if not patterns:
        return text
    
    cleaned_text = text
    for pattern_info in patterns:
        try:
            regex_pattern = pattern_info.get("regex", "")
            if regex_pattern:
                # Компилируем паттерн
                compiled_pattern = re.compile(regex_pattern)
                # Удаляем совпадения
                cleaned_text = compiled_pattern.sub("", cleaned_text)
        except Exception as e:
            # Если паттерн некорректный, пропускаем его
            print(f"[WARNING] Ошибка при применении паттерна '{pattern_info.get('name', 'unknown')}': {e}")
            continue
    
    # Удаляем множественные пустые строки (более 2 подряд)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    # Удаляем пробелы в начале и конце
    cleaned_text = cleaned_text.strip()
    
    return cleaned_text

def clean_tree_with_patterns(node: Dict[str, Any], patterns: List[Dict[str, str]] = None):
    """Рекурсивно очищает content всех узлов дерева используя паттерны"""
    if patterns is None:
        patterns = load_cleaning_patterns()
    
    # Очищаем content текущего узла
    if "content" in node and node["content"]:
        node["content"] = clean_text_with_patterns(node["content"], patterns)
    
    # Рекурсивно обрабатываем children
    if "children" in node and node["children"]:
        for child in node["children"]:
            clean_tree_with_patterns(child, patterns)

# ---------- 3) Система обработчиков документов ----------

class DocumentProcessor(ABC):
    """Базовый класс для обработчиков документов"""
    
    @abstractmethod
    def can_process(self, path: str) -> bool:
        """Проверяет, может ли обработчик обработать данный документ"""
        pass
    
    @abstractmethod
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ и возвращает дерево JSON"""
        pass
    
    def convert_to_markdown(self, path: str) -> str:
        """Универсальный метод преобразования в markdown через markitdown"""
        md = MarkItDown(enable_plugins=False)
        res = md.convert(path)
        return (res.text_content or "").strip()
    
    def apply_cleaning_patterns(self, tree: Dict[str, Any]) -> Dict[str, Any]:
        """Применяет паттерны очистки к дереву документа"""
        clean_tree_with_patterns(tree)
        return tree

# ---------- 3) Обработчик для документов Гаранта (ГАРАНТ) ----------

class GarantProcessor(DocumentProcessor):
    """Обработчик для документов системы ГАРАНТ"""
    
    def can_process(self, path: str) -> bool:
        """Определяет документы Гаранта по характерным признакам"""
        try:
            # Проверяем содержимое документа
            md = self.convert_to_markdown(path)
            # Характерные признаки документов Гаранта:
            # - наличие "ГАРАНТ:" в тексте
            # - наличие ссылок на ivo.garant.ru
            # - структура "Глава X", "Статья X"
            has_garant_marker = "ГАРАНТ:" in md or "ГАРАНТ" in md
            has_garant_links = "ivo.garant.ru" in md or "garant.ru" in md
            has_legal_structure = bool(re.search(r'Глава\s+\d+|Статья\s+\d+', md))
            return has_garant_marker or (has_garant_links and has_legal_structure)
        except:
            return False
    
    def _clean_garant_text(self, text: str) -> str:
        """Удаляет служебные строки ГАРАНТ из текста"""
        lines = text.splitlines()
        cleaned_lines = []
        
        # Паттерны для удаления
        garant_pattern = re.compile(r'^\s*ГАРАНТ\s*:?\s*$', re.IGNORECASE)
        # Паттерн для "См. комментарии к статье X" (с возможными вариантами окончания)
        comment_pattern = re.compile(r'^\s*См\.\s+комментарии\s+к\s+статье\s+\d+.*?$', re.IGNORECASE)
        
        i = 0
        while i < len(lines):
            line = lines[i]
            # Пропускаем строки "ГАРАНТ:" или "ГАРАНТ"
            if garant_pattern.match(line):
                i += 1
                # Пропускаем пустые строки после "ГАРАНТ:"
                while i < len(lines) and not lines[i].strip():
                    i += 1
                # Пропускаем строку "См. комментарии ..." если она идет после "ГАРАНТ:"
                if i < len(lines) and comment_pattern.match(lines[i]):
                    i += 1
                    # Пропускаем пустые строки после комментария
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                continue
            
            # Пропускаем строки "См. комментарии ..." отдельно
            if comment_pattern.match(line):
                i += 1
                # Пропускаем пустые строки после комментария
                while i < len(lines) and not lines[i].strip():
                    i += 1
                continue
            
            cleaned_lines.append(line)
            i += 1
        
        return "\n".join(cleaned_lines)
    
    def _clean_tree_content(self, node: Dict[str,Any]):
        """Рекурсивно очищает контент узлов дерева от служебных строк ГАРАНТ"""
        if "content" in node and node["content"]:
            node["content"] = self._clean_garant_text(node["content"])
        for child in node.get("children", []):
            self._clean_tree_content(child)
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ Гаранта"""
        md = self.convert_to_markdown(path)
        
        # Очищаем от служебных строк ГАРАНТ перед обработкой
        md = self._clean_garant_text(md)
        
        # Дополнительная обработка для документов Гаранта:
        # - Улучшаем распознавание заголовков глав и статей
        # - Сохраняем ссылки в формате [текст](url)
        
        # Паттерны для глав и статей
        chapter_pattern = re.compile(r'^(Глава\s+\d+[\.:]?\s*.+?)$', re.MULTILINE | re.IGNORECASE)
        article_pattern = re.compile(r'^(Статья\s+\d+[\.:]?\s*.+?)$', re.MULTILINE | re.IGNORECASE)
        
        # Преобразуем в заголовки markdown
        lines = md.splitlines()
        processed_lines = []
        i = 0
        while i < len(lines):
            line = lines[i]
            # Проверяем на главу
            if chapter_pattern.match(line.strip()):
                processed_lines.append(f"# {line.strip()}")
                i += 1
                continue
            # Проверяем на статью
            if article_pattern.match(line.strip()):
                processed_lines.append(f"## {line.strip()}")
                i += 1
                continue
            processed_lines.append(line)
            i += 1
        
        md_processed = "\n".join(processed_lines)
        tree = parse_markdown_to_tree(md_processed)
        
        # Дополнительная очистка уже распарсенного дерева (на случай если что-то пропустили)
        self._clean_tree_content(tree)
        
        # Применяем паттерны очистки из JSON файла
        tree = self.apply_cleaning_patterns(tree)
        
        return tree

# ---------- 4) Обработчик для документов Консультанта (КонсультантПлюс) ----------

class ConsultantProcessor(DocumentProcessor):
    """Обработчик для документов системы КонсультантПлюс"""
    
    def can_process(self, path: str) -> bool:
        """Определяет документы Консультанта по характерным признакам"""
        try:
            md = self.convert_to_markdown(path)
            # Характерные признаки документов Консультанта:
            # - наличие ссылок на consultant.ru
            # - наличие "КонсультантПлюс" в тексте
            # - структура с разделами, главами, статьями и подразделами
            has_consultant_links = "consultant.ru" in md or "КонсультантПлюс" in md or "Консультант" in md
            has_sections = bool(re.search(r'Раздел\s+\d+|Подраздел\s+\d+|Глава\s+\d+|Статья\s+\d+', md, re.IGNORECASE))
            return has_consultant_links or has_sections
        except:
            return False
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ Консультанта"""
        md = self.convert_to_markdown(path)
        
        # Дополнительная обработка для документов Консультанта:
        # - Улучшаем распознавание разделов с римскими цифрами
        # - Распознавание подзаголовков
        # - Распознавание разделов и подразделов
        
        lines = md.splitlines()
        processed_lines = []
        i = 0
        
        # Паттерны для распознавания заголовков
        section_pattern = re.compile(r'^(Раздел\s+\d+[\.:]?\s*.+)$', re.IGNORECASE)
        subsection_pattern = re.compile(r'^(Подраздел\s+\d+[\.:]?\s*.+)$', re.IGNORECASE)
        # Глава с номером: "Глава 1", "Глава 2" и т.д. (может быть с точкой или без)
        chapter_pattern = re.compile(r'^(Глава\s+\d+[\.:]?\s*.+)$', re.IGNORECASE)
        # Статья с номером: "Статья 1", "Статья 2" и т.д. (может быть с точкой или без)
        article_pattern = re.compile(r'^(Статья\s+\d+[\.:]?\s*.+)$', re.IGNORECASE)
        # Римские цифры с точкой: I., II., III., IV., V.
        roman_section_pattern = re.compile(r'^([IVXLCDM]+)\.\s+(.+)$', re.IGNORECASE)
        # Подзаголовки (начинаются с заглавной буквы, короткие, без точки в конце)
        subtitle_pattern = re.compile(r'^([А-ЯЁ][А-Яа-яё\s]{5,80})$')
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Пропускаем пустые строки
            if not line:
                processed_lines.append("")
                i += 1
                continue
            
            # Пропускаем строки таблиц (начинаются с |)
            if line.startswith('|'):
                processed_lines.append(lines[i])
                i += 1
                continue
            
            # Проверяем на раздел с римской цифрой
            roman_match = roman_section_pattern.match(line)
            if roman_match:
                section_title = roman_match.group(2).strip()
                processed_lines.append(f"# {line}")
                i += 1
                continue
            
            # Проверяем на Главу (может быть на нескольких строках)
            if chapter_pattern.match(line):
                chapter_text = line
                # Собираем следующие строки, если они являются продолжением заголовка главы
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    # Если следующая строка пустая или начинается с "Статья" или "Глава" - прекращаем
                    if not next_line or article_pattern.match(next_line) or chapter_pattern.match(next_line):
                        break
                    # Если следующая строка - таблица или другой заголовок - прекращаем
                    if next_line.startswith('|') or section_pattern.match(next_line) or subsection_pattern.match(next_line):
                        break
                    # Если следующая строка начинается с заглавной буквы и короткая - возможно продолжение
                    if len(next_line) < 100 and next_line.isupper():
                        chapter_text += " " + next_line
                        j += 1
                    else:
                        break
                processed_lines.append(f"# {chapter_text}")
                i = j
                continue
            
            # Проверяем на Статью
            if article_pattern.match(line):
                processed_lines.append(f"## {line}")
                i += 1
                continue
            
            # Проверяем на обычный раздел
            if section_pattern.match(line):
                processed_lines.append(f"# {line}")
                i += 1
                continue
            
            # Проверяем на подраздел
            if subsection_pattern.match(line):
                processed_lines.append(f"## {line}")
                i += 1
                continue
            
            # Проверяем на подзаголовок (короткая строка, начинается с заглавной, без точки)
            # Но не пронумерованный пункт
            if (subtitle_pattern.match(line) and 
                not _is_numbered_item(line) and 
                len(line) < 100 and
                not line.endswith('.') and
                not line.endswith(',') and
                i + 1 < len(lines) and  # Есть следующая строка
                lines[i + 1].strip()):  # Следующая строка не пустая
                # Проверяем, что следующая строка не является заголовком
                next_line = lines[i + 1].strip()
                if not (roman_section_pattern.match(next_line) or 
                       section_pattern.match(next_line) or
                       subsection_pattern.match(next_line) or
                       chapter_pattern.match(next_line) or
                       article_pattern.match(next_line)):
                    processed_lines.append(f"## {line}")
                    i += 1
                    continue
            
            processed_lines.append(lines[i])
            i += 1
        
        md_processed = "\n".join(processed_lines)
        tree = parse_markdown_to_tree(md_processed)
        return self.apply_cleaning_patterns(tree)

# ---------- 5) Обработчик для DOCX со стилями заголовков ----------

class DocxWithHeadingStylesProcessor(DocumentProcessor):
    """Обработчик для DOCX со стилями заголовков (Heading 1, Heading 2, etc.)"""
    
    def can_process(self, path: str) -> bool:
        """Проверяет наличие стилей заголовков в DOCX"""
        if not path.lower().endswith('.docx'):
            return False
        try:
            doc = Document(path)
            # Проверяем наличие стилей заголовков
            heading_styles = set()
            for p in doc.paragraphs:
                style_name = p.style.name if p.style else ""
                if re.match(r'Heading\s+\d+', style_name, re.IGNORECASE):
                    heading_styles.add(style_name)
            return len(heading_styles) > 0
        except:
            return False
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает DOCX со стилями заголовков через markitdown"""
        md = self.convert_to_markdown(path)
        # markitdown должен правильно обработать стили заголовков
        tree = parse_markdown_to_tree(md)
        return self.apply_cleaning_patterns(tree)

# ---------- 6) Обработчик для DOCX со стилями, не наследуемыми от заголовков ----------

class DocxWithCustomStylesProcessor(DocumentProcessor):
    """Обработчик для DOCX с кастомными стилями, не наследуемыми от заголовков"""
    
    def can_process(self, path: str) -> bool:
        """Проверяет наличие кастомных стилей (не Heading)"""
        if not path.lower().endswith('.docx'):
            return False
        try:
            doc = Document(path)
            # Проверяем наличие outline levels или кастомных стилей
            has_outline = False
            has_custom_styles = False
            for p in doc.paragraphs:
                # Проверяем outline level
                el = p._element
                outline_vals = el.xpath(".//w:outlineLvl/@w:val")
                if outline_vals:
                    has_outline = True
                # Проверяем кастомные стили (не Heading и не Normal)
                style_name = (p.style.name or "").strip()
                if style_name and not re.match(r'Heading\s+\d+|Normal', style_name, re.IGNORECASE):
                    has_custom_styles = True
            return has_outline or has_custom_styles
        except:
            return False
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает DOCX с кастомными стилями через fallback метод"""
        # Используем fallback метод, который анализирует outline levels и типографику
        md = self._docx_fallback_to_markdown(path)
        tree = parse_markdown_to_tree(md)
        return self.apply_cleaning_patterns(tree)
    
    def _docx_fallback_to_markdown(self, path: str) -> str:
        """Превращает DOCX в псевдо-Markdown используя outline levels и типографику"""
        doc = Document(path)
        infos = self._gather_paragraph_info(doc)
        lines = []
        for is_h, lvl, text in self._score_and_level_candidates(infos):
            if is_h:
                lines.append("#"*lvl + " " + text)
            else:
                lines.append(text)
            lines.append("")  # пустая строка-разделитель
        return "\n".join(lines).strip()
    
    def _gather_paragraph_info(self, doc: Document) -> List[Any]:
        """Собирает информацию о параграфах"""
        from dataclasses import dataclass
        
        @dataclass
        class PInfo:
            text: str
            lvl_heading_style: Optional[int]
            lvl_outline: Optional[int]
            lvl_list: Optional[int]
            font_pt: Optional[float]
            bold_ratio: float
            upper_ratio: float
            is_short: bool
        
        infos: List[PInfo] = []
        for p in doc.paragraphs:
            text, runs = self._run_text_and_attrs(p)
            if not text:
                continue
            bold_chars = sum(len(t) for (t,b,_) in runs if b)
            total_chars = sum(len(t) for (t,_,_) in runs)
            bold_ratio = (bold_chars / total_chars) if total_chars else 0.0
            
            letters = re.findall(r"[A-Za-zА-Яа-яЁё]", text)
            uppers = [c for c in letters if c.isupper()]
            upper_ratio = (len(uppers) / len(letters)) if letters else 0.0
            
            font_pts = [sz for (_,_,sz) in runs if sz is not None]
            font_pt = max(font_pts) if font_pts else None
            
            infos.append(PInfo(
                text=text,
                lvl_heading_style=self._heading_style_level(p),
                lvl_outline=self._outline_level(p),
                lvl_list=self._list_ilvl(p),
                font_pt=font_pt,
                bold_ratio=bold_ratio,
                upper_ratio=upper_ratio,
                is_short=(len(text) <= 120)
            ))
        return infos
    
    def _run_text_and_attrs(self, p) -> Tuple[str, List[Tuple[str, bool, Optional[float]]]]:
        txt = []
        runs = []
        for r in p.runs:
            t = r.text or ""
            if not t: continue
            txt.append(t)
            sz = r.font.size.pt if r.font.size is not None else None
            runs.append((t, bool(r.bold), sz))
        return "".join(txt).strip(), runs
    
    def _outline_level(self, p) -> Optional[int]:
        el = p._element
        vals = el.xpath(".//w:outlineLvl/@w:val")
        if vals:
            try:
                return int(vals[0]) + 1
            except Exception:
                return None
        return None
    
    def _list_ilvl(self, p) -> Optional[int]:
        el = p._element
        vals = el.xpath(".//w:numPr/w:ilvl/@w:val")
        if vals:
            try:
                return int(vals[0])
            except Exception:
                return None
        return None
    
    def _heading_style_level(self, p) -> Optional[int]:
        try:
            name = (p.style.name or "").strip()
        except Exception:
            return None
        m = re.fullmatch(r"Heading\s+([1-9])", name, flags=re.IGNORECASE)
        return int(m.group(1)) if m else None
    
    def _rank_font_sizes(self, infos: List[Any]) -> Dict[float,int]:
        sizes = sorted({i.font_pt for i in infos if i.font_pt is not None}, reverse=True)
        if not sizes:
            return {}
        return {sz: lvl for lvl, sz in enumerate(sizes[:6], start=1)}
    
    def _calibrate_list_base(self, infos: List[Any]) -> int:
        lvls = [i.lvl_list for i in infos if i.lvl_list is not None]
        if not lvls:
            return 1
        return 1 - min(lvls)
    
    def _score_and_level_candidates(self, infos: List[Any]) -> List[Tuple[bool, int, str]]:
        font_rank = self._rank_font_sizes(infos)
        base = self._calibrate_list_base(infos)
        
        out: List[Tuple[bool,int,str]] = []
        for i in infos:
            cand_levels = []
            
            if i.lvl_outline is not None:
                cand_levels.append(("outline", i.lvl_outline, 1.0))
            if i.lvl_heading_style is not None:
                cand_levels.append(("hstyle", i.lvl_heading_style, 0.95))
            if i.lvl_list is not None:
                cand_levels.append(("list", max(1, min(6, i.lvl_list + base)), 0.8))
            
            if i.font_pt is not None and font_rank:
                if i.font_pt in font_rank:
                    cand_levels.append(("font", font_rank[i.font_pt], 0.7))
            if i.is_short and (i.bold_ratio >= 0.6 or i.upper_ratio >= 0.6):
                cand_levels.append(("typo", 3, 0.55))
            
            if cand_levels:
                best = max(cand_levels, key=lambda t: t[2])
                lvl = max(1, min(6, best[1]))
                out.append((True, lvl, i.text))
            else:
                out.append((False, 0, i.text))
        return out

# ---------- 7) Обработчик для ГОСТов ----------

class GOSTProcessor(DocumentProcessor):
    """Обработчик для документов ГОСТ (ГОСТы)"""
    
    def can_process(self, path: str) -> bool:
        """Определяет документы ГОСТ по характерным признакам"""
        try:
            md = self.convert_to_markdown(path)
            # Характерные признаки ГОСТов:
            # - наличие "ГОСТ" в начале документа
            # - наличие "МЕЖГОСУДАРСТВЕННЫЙ СТАНДАРТ" или "ГОСУДАРСТВЕННЫЙ СТАНДАРТ"
            # - структура с разделами типа "1. ОБЩИЕ ПОЛОЖЕНИЯ", "2. СОСТАВ И СОДЕРЖАНИЕ"
            has_gost = bool(re.search(r'^ГОСТ\s+\d+', md, re.MULTILINE))
            has_standard = "МЕЖГОСУДАРСТВЕННЫЙ СТАНДАРТ" in md or "ГОСУДАРСТВЕННЫЙ СТАНДАРТ" in md
            has_gost_structure = bool(re.search(r'^\d+\.\s+[А-ЯЁ\s]+$', md, re.MULTILINE))
            return has_gost or (has_standard and has_gost_structure)
        except:
            return False
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ ГОСТ"""
        md = self.convert_to_markdown(path)
        
        lines = md.splitlines()
        processed_lines = []
        i = 0
        
        # Паттерны для распознавания структуры ГОСТа
        # Разделы: "1. ОБЩИЕ ПОЛОЖЕНИЯ", "2. СОСТАВ И СОДЕРЖАНИЕ"
        section_pattern = re.compile(r'^(\d+)\.\s+([А-ЯЁ\s]+)$')
        # Подразделы: "1.1. Текст", "2.3. Текст", "2.4.1. Текст" (многоуровневые)
        subsection_pattern = re.compile(r'^(\d+)(?:\.(\d+))+(?:\.\s+(.+))?$')
        # Пункты: "1) текст", "2) текст"
        item_pattern = re.compile(r'^(\d+)\)\s+(.+)$')
        # Подпункты: "а) текст", "б) текст"
        subitem_pattern = re.compile(r'^([а-яА-Я])\)\s+(.+)$')
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Пропускаем пустые строки
            if not line:
                processed_lines.append("")
                i += 1
                continue
            
            # Пропускаем строки таблиц
            if line.startswith('|'):
                processed_lines.append(lines[i])
                i += 1
                continue
            
            # Проверяем на раздел (например, "1. ОБЩИЕ ПОЛОЖЕНИЯ")
            section_match = section_pattern.match(line)
            if section_match:
                processed_lines.append(f"# {line}")
                i += 1
                continue
            
            # Проверяем на подраздел (например, "1.1. Текст", "2.4.1. Текст")
            # Сначала проверяем многоуровневые подразделы (2.4.1., 2.4.2. и т.д.)
            if re.match(r'^\d+\.\d+\.\d+\.', line):
                # Подраздел третьего уровня
                processed_lines.append(f"### {line}")
                i += 1
                continue
            elif re.match(r'^\d+\.\d+\.', line):
                # Подраздел второго уровня
                processed_lines.append(f"## {line}")
                i += 1
                continue
            
            # Проверяем на пункт (например, "1) текст")
            # Но сначала проверяем, нет ли в строке нескольких пунктов через точку с запятой
            if ';' in line and re.search(r'\d+\)\s+', line):
                # Если в строке несколько пунктов, разделяем их
                # Ищем все пункты в строке
                items_in_line = re.findall(r'(\d+\)\s+[^;]+(?:;|$))', line)
                if len(items_in_line) > 1:
                    # Разделяем на отдельные пункты
                    for item in items_in_line:
                        item_clean = item.strip()
                        if item_clean:
                            processed_lines.append(f"#### {item_clean}")
                    i += 1
                    continue
            
            item_match = item_pattern.match(line)
            if item_match:
                processed_lines.append(f"#### {line}")
                i += 1
                continue
            
            # Проверяем на подпункт (например, "а) текст")
            subitem_match = subitem_pattern.match(line)
            if subitem_match:
                processed_lines.append(f"##### {line}")
                i += 1
                continue
            
            # Обычный текст
            processed_lines.append(lines[i])
            i += 1
        
        md_processed = "\n".join(processed_lines)
        tree = parse_markdown_to_tree(md_processed)
        return self.apply_cleaning_patterns(tree)

# ---------- 8) Универсальный обработчик для прочих документов ----------

class UniversalProcessor(DocumentProcessor):
    """Универсальный обработчик для всех остальных документов"""
    
    def can_process(self, path: str) -> bool:
        """Всегда может обработать (fallback)"""
        return True
    
    def process(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ через markitdown"""
        md = self.convert_to_markdown(path)
        tree = parse_markdown_to_tree(md)
        return self.apply_cleaning_patterns(tree)

# ---------- 9) Менеджер обработчиков ----------

class DocumentProcessorManager:
    """Менеджер для выбора и использования обработчиков"""
    
    def __init__(self):
        self.processors: List[DocumentProcessor] = [
            GarantProcessor(),
            ConsultantProcessor(),
            GOSTProcessor(),
            DocxWithHeadingStylesProcessor(),
            DocxWithCustomStylesProcessor(),
            UniversalProcessor()  # Всегда последний как fallback
        ]
    
    def process_document(self, path: str) -> Dict[str,Any]:
        """Обрабатывает документ используя подходящий обработчик"""
        for processor in self.processors:
            if processor.can_process(path):
                print(f"[INFO] Используется обработчик: {processor.__class__.__name__}")
                return processor.process(path)
        # Fallback на универсальный
        return UniversalProcessor().process(path)

# ---------- 9) Утилиты сохранения/загрузки ----------

REQUIRED_KEYS = {"level","title","content","children"}

def _validate_tree(node: Dict[str,Any], path="root"):
    if not isinstance(node, dict): raise TypeError(f"{path}: node must be dict")
    missing = REQUIRED_KEYS - set(node.keys())
    if missing: raise ValueError(f"{path}: missing keys {missing}")
    if not isinstance(node["level"], int): raise TypeError(f"{path}.level must be int")
    if not isinstance(node["title"], str): raise TypeError(f"{path}.title must be str")
    if not isinstance(node["content"], str): raise TypeError(f"{path}.content must be str")
    if not isinstance(node["children"], list): raise TypeError(f"{path}.children must be list")
    for i,ch in enumerate(node["children"]): _validate_tree(ch, f"{path}.children[{i}]")

def save_tree_json(tree: Dict[str,Any], out_path: str, pretty=True):
    _validate_tree(tree)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    if pretty:
        txt = json.dumps(tree, ensure_ascii=False, indent=2)
    else:
        txt = json.dumps(tree, ensure_ascii=False, separators=(",",":"))
    Path(out_path).write_text(txt, encoding="utf-8")

# ---------- 10) Публичный API ----------

def extract_outline_from_document(path: str) -> Dict[str,Any]:
    """Универсальная функция для извлечения структуры из документа"""
    manager = DocumentProcessorManager()
    return manager.process_document(path)

# ---------- 11) CLI-пример ----------

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Convert documents to JSON outline (sections hierarchy).")
    ap.add_argument("inputs", nargs="+", help="Paths to documents (DOCX, PDF, etc.)")
    ap.add_argument("-o","--out", default="out", help="Output folder (default: ./out)")
    args = ap.parse_args()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    manager = DocumentProcessorManager()
    
    for p in args.inputs:
        try:
            tree = manager.process_document(p)
            save_tree_json(tree, out_dir / (Path(p).stem + ".json"), pretty=True)
            print(f"[ok] {p} → {out_dir / (Path(p).stem + '.json')}")
        except Exception as e:
            print(f"[error] {p}: {e}")
